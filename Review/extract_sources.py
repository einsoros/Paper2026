from pathlib import Path
import re

from docx import Document
from openpyxl import load_workbook


SOURCE_DIR = Path("/Users/user/Library/Mobile Documents/com~apple~CloudDocs/Desktop/Paper/1. 논문작성")
OUT_DIR = Path("Paper/Review/source_extracts")

FILES = [
    "연구 구조 요약서_260522.docx",
    "논문 구조 및 지지이론_260519.xlsx",
    "논문 목차_260512.xlsx",
    "논문_구조 및 키워드.docx",
]


def slug(name: str) -> str:
    name = re.sub(r"\.(docx|xlsx)$", "", name)
    name = re.sub(r"[^\w가-힣]+", "_", name)
    return name.strip("_")


def clean(value):
    if value is None:
        return ""
    return str(value).replace("\n", " ").strip()


def extract_docx(path: Path) -> str:
    doc = Document(path)
    lines = [f"# {path.name}", ""]

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style else ""
            if style.startswith("Heading"):
                lines.append(f"## {text}")
            else:
                lines.append(text)
            lines.append("")

    for idx, table in enumerate(doc.tables, start=1):
        lines.append(f"## Table {idx}")
        lines.append("")
        for row in table.rows:
            cells = [clean(cell.text) for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def extract_xlsx(path: Path) -> str:
    wb = load_workbook(path, data_only=False)
    lines = [f"# {path.name}", ""]

    for ws in wb.worksheets:
        lines.append(f"## Sheet: {ws.title}")
        lines.append("")
        max_row = ws.max_row
        max_col = ws.max_column
        for row in ws.iter_rows(min_row=1, max_row=max_row, max_col=max_col):
            cells = [clean(cell.value) for cell in row]
            while cells and cells[-1] == "":
                cells.pop()
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    index_lines = ["# Source Extracts", ""]

    for file_name in FILES:
        src = SOURCE_DIR / file_name
        out = OUT_DIR / f"{slug(file_name)}.md"
        if file_name.endswith(".docx"):
            content = extract_docx(src)
        elif file_name.endswith(".xlsx"):
            content = extract_xlsx(src)
        else:
            continue
        out.write_text(content, encoding="utf-8")
        index_lines.append(f"- [{file_name}]({out.name})")

    (OUT_DIR / "index.md").write_text("\n".join(index_lines) + "\n", encoding="utf-8")


if __name__ == "__main__":
    main()
